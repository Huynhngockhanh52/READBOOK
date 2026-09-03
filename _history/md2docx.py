#!/usr/bin/env python3
"""Convert Markdown files to DOCX files based on a Word template."""

from __future__ import annotations

import argparse
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Sequence


dep_err: Exception | None = None
try:
    import mistune
    from docx import Document as DocxDoc
    from docx.enum.style import WD_STYLE_TYPE as WdStyle
    from docx.shared import Inches as toInches
except ImportError as exc:
    dep_err = exc


style_map = {
    1: "Heading 1",
    2: "KC H2",
    3: "KC H3",
    4: "KC H4",
}
norm_style = "KC NORMAL"
req_styles = tuple(style_map.values()) + (norm_style,)


class ConvertError(RuntimeError):
    """Raised when Markdown cannot be converted safely."""


@dataclass(frozen=True)
class FileJob:
    src_path: Path
    out_path: Path


def getRoot() -> Path:
    return Path(__file__).resolve().parent.parent


def needDeps() -> None:
    if dep_err is not None:
        raise ConvertError(
            "Thiếu dependency. Hãy chạy: pip install python-docx mistune"
        ) from dep_err


def parseArgs(argv: Sequence[str] | None = None) -> argparse.Namespace:
    root = getRoot()
    parser = argparse.ArgumentParser(
        description="Chuyển Markdown sang DOCX dựa trên Word template.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "-i",
        "--input",
        dest="input_path",
        type=Path,
        required=True,
        help="Một file .md hoặc thư mục chứa các file .md.",
    )
    parser.add_argument(
        "-o",
        "--output",
        dest="output_path",
        type=Path,
        help="Thư mục output DOCX.",
    )
    parser.add_argument(
        "--template",
        type=Path,
        default=root / "sample" / "sample.docx",
        help="Word template chứa các style bắt buộc.",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Tạo lại file DOCX đã tồn tại.",
    )
    return parser.parse_args(argv)


def getMdFiles(path: Path) -> tuple[Path, list[Path]]:
    in_path = path.expanduser().resolve()
    if not in_path.exists():
        raise ConvertError(f"Input không tồn tại: {in_path}")
    if in_path.is_file():
        if in_path.suffix.lower() != ".md":
            raise ConvertError(f"Input phải là file .md: {in_path}")
        return in_path, [in_path]
    if not in_path.is_dir():
        raise ConvertError(f"Input không phải file hoặc thư mục: {in_path}")

    files = sorted(
        (p for p in in_path.glob("*.md") if p.is_file()),
        key=lambda p: p.name.lower(),
    )
    if not files:
        raise ConvertError(f"Không có file .md trong thư mục: {in_path}")
    return in_path, files


def getOutDir(in_path: Path, out_arg: Path | None) -> Path:
    if out_arg is not None:
        out_dir = out_arg.expanduser().resolve()
        if out_dir.exists() and not out_dir.is_dir():
            raise ConvertError(f"Output phải là thư mục: {out_dir}")
        return out_dir
    if in_path.is_file():
        return in_path.parent.parent / "docx"
    return in_path.parent / "docx"


def getOutName(src_path: Path) -> str:
    stem = src_path.stem
    if stem.lower().endswith("_vi"):
        stem = stem[:-3] or src_path.stem
    return f"{stem}.docx"


def makeJobs(files: list[Path], out_dir: Path) -> list[FileJob]:
    return [FileJob(p, out_dir / getOutName(p)) for p in files]


def checkTpl(tpl_path: Path) -> None:
    if not tpl_path.is_file():
        raise ConvertError(f"Không tìm thấy template: {tpl_path}")
    try:
        doc = DocxDoc(tpl_path)
    except Exception as exc:
        raise ConvertError(f"Không thể mở template {tpl_path}: {exc}") from exc

    missing: list[str] = []
    for name in req_styles:
        try:
            style = doc.styles[name]
        except KeyError:
            missing.append(name)
            continue
        if style.type != WdStyle.PARAGRAPH:
            missing.append(f"{name} (không phải paragraph style)")
    if missing:
        names = ", ".join(missing)
        raise ConvertError(f"Template thiếu style bắt buộc: {names}")


def makeAst(text: str) -> list[dict[str, Any]]:
    try:
        parser = mistune.create_markdown(renderer="ast")
        nodes = parser(text)
    except Exception as exc:
        raise ConvertError(f"Không thể parse Markdown: {exc}") from exc
    if not isinstance(nodes, list):
        raise ConvertError("Markdown parser trả về cấu trúc không hợp lệ.")
    return nodes


def clearBody(doc: Any) -> None:
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


def nodeText(node: dict[str, Any]) -> str:
    raw = node.get("raw")
    if isinstance(raw, str):
        return raw
    kids = node.get("children") or []
    return "".join(nodeText(n) for n in kids if isinstance(n, dict))


def addRun(
    para: Any,
    text: str,
    is_bold: bool = False,
    is_italic: bool = False,
    is_code: bool = False,
) -> Any:
    run = para.add_run(text)
    if is_bold:
        run.bold = True
    if is_italic:
        run.italic = True
    if is_code:
        run.font.name = "Consolas"
    return run


def addInline(
    para: Any,
    nodes: list[dict[str, Any]],
    is_bold: bool = False,
    is_italic: bool = False,
) -> None:
    for node in nodes:
        node_type = str(node.get("type", ""))
        kids = node.get("children") or []
        if node_type == "text":
            addRun(para, str(node.get("raw", "")), is_bold, is_italic)
        elif node_type == "strong":
            addInline(para, kids, True, is_italic)
        elif node_type == "emphasis":
            addInline(para, kids, is_bold, True)
        elif node_type == "codespan":
            addRun(
                para,
                str(node.get("raw", "")),
                is_bold,
                is_italic,
                True,
            )
        elif node_type in {"linebreak", "softbreak"}:
            addRun(para, "\n", is_bold, is_italic)
        elif node_type == "link":
            addInline(para, kids, is_bold, is_italic)
            url = str((node.get("attrs") or {}).get("url", ""))
            label = nodeText(node)
            if url and url not in label:
                addRun(para, f" ({url})", is_bold, is_italic)
        elif node_type == "image":
            label = nodeText(node) or "image"
            url = str((node.get("attrs") or {}).get("url", ""))
            value = f"[{label}]"
            if url:
                value += f" ({url})"
            addRun(para, value, is_bold, is_italic)
        elif isinstance(node.get("raw"), str):
            addRun(para, str(node["raw"]), is_bold, is_italic)
        elif kids:
            addInline(para, kids, is_bold, is_italic)


def setIndent(para: Any, depth: int, hanging: bool = False) -> None:
    fmt = para.paragraph_format
    fmt.left_indent = toInches(0.25 * (depth + 1))
    if hanging:
        fmt.first_line_indent = toInches(-0.18)


def addPara(
    doc: Any,
    nodes: list[dict[str, Any]],
    style: str = norm_style,
    prefix: str = "",
    depth: int = -1,
    quote: bool = False,
) -> Any:
    para = doc.add_paragraph(style=style)
    if depth >= 0:
        setIndent(para, depth, True)
    elif quote:
        setIndent(para, 0)
    if prefix:
        addRun(para, prefix)
    addInline(para, nodes)
    return para


def addList(doc: Any, node: dict[str, Any], depth: int = 0) -> None:
    attrs = node.get("attrs") or {}
    ordered = bool(attrs.get("ordered", False))
    start_num = int(attrs.get("start") or 1)
    items = node.get("children") or []

    for i, item in enumerate(items):
        prefix = f"{start_num + i}. " if ordered else "• "
        first = True
        kids = item.get("children") or []
        for child in kids:
            kind = str(child.get("type", ""))
            if kind in {"block_text", "paragraph"}:
                mark = prefix if first else ""
                addPara(
                    doc,
                    child.get("children") or [],
                    prefix=mark,
                    depth=depth,
                )
                first = False
            elif kind == "list":
                addList(doc, child, depth + 1)
            else:
                addBlock(doc, child, depth)
        if first:
            addPara(doc, [], prefix=prefix, depth=depth)


def addBlock(
    doc: Any,
    node: dict[str, Any],
    depth: int = 0,
    quote: bool = False,
) -> None:
    kind = str(node.get("type", ""))
    kids = node.get("children") or []

    if kind == "heading":
        level = int((node.get("attrs") or {}).get("level", 1))
        style = style_map.get(level, norm_style)
        addPara(doc, kids, style=style)
    elif kind in {"paragraph", "block_text"}:
        addPara(doc, kids, quote=quote)
    elif kind == "block_code":
        para = doc.add_paragraph(style=norm_style)
        if quote:
            setIndent(para, 0)
        run = addRun(para, str(node.get("raw", "")).rstrip("\n"), is_code=True)
        run.font.name = "Consolas"
        para.paragraph_format.keep_together = True
    elif kind == "list":
        addList(doc, node, depth)
    elif kind == "block_quote":
        for child in kids:
            addBlock(doc, child, depth, True)
    elif kind == "thematic_break":
        addPara(doc, [{"type": "text", "raw": "────────────────"}])
    elif kind == "blank_line":
        return
    elif isinstance(node.get("raw"), str):
        addPara(doc, [{"type": "text", "raw": str(node["raw"])}], quote=quote)
    elif kids:
        for child in kids:
            addBlock(doc, child, depth, quote)


def addBlocks(doc: Any, nodes: list[dict[str, Any]]) -> None:
    for node in nodes:
        addBlock(doc, node)


def makeDocx(src_path: Path, out_path: Path, tpl_path: Path) -> None:
    if out_path.resolve() == tpl_path.resolve():
        raise ConvertError("Output không được trùng với file template.")
    try:
        text = src_path.read_text(encoding="utf-8")
    except (OSError, UnicodeError) as exc:
        raise ConvertError(f"Không thể đọc Markdown {src_path}: {exc}") from exc

    nodes = makeAst(text)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    copy_made = False
    try:
        shutil.copy2(tpl_path, out_path)
        copy_made = True
        doc = DocxDoc(out_path)
        clearBody(doc)
        addBlocks(doc, nodes)
        doc.save(out_path)
        DocxDoc(out_path)
    except Exception as exc:
        if copy_made and out_path.exists():
            try:
                out_path.unlink()
            except OSError:
                pass
        if isinstance(exc, ConvertError):
            raise
        raise ConvertError(f"Không thể tạo DOCX {out_path}: {exc}") from exc


def main(argv: Sequence[str] | None = None) -> int:
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8", errors="replace")

    args = parseArgs(argv)
    try:
        needDeps()
        in_path, files = getMdFiles(args.input_path)
        out_dir = getOutDir(in_path, args.output_path)
        tpl_path = args.template.expanduser().resolve()
        checkTpl(tpl_path)
        jobs = makeJobs(files, out_dir)

        done = 0
        skipped = 0
        failed: list[tuple[Path, str]] = []
        total = len(jobs)
        for i, job in enumerate(jobs, start=1):
            print(f"[{i}/{total}] Processing {job.src_path.name}")
            if job.out_path.exists() and not args.overwrite:
                print(f"[SKIP] Output already exists: {job.out_path}")
                skipped += 1
                continue
            try:
                makeDocx(job.src_path, job.out_path, tpl_path)
                print(f"[DONE] {job.out_path}")
                done += 1
            except ConvertError as exc:
                print(f"[ERROR] {job.src_path}: {exc}", file=sys.stderr)
                failed.append((job.src_path, str(exc)))

        print(
            f"Completed: files={total}, converted={done}, "
            f"skipped={skipped}, failed={len(failed)}"
        )
        return 1 if failed else 0
    except ConvertError as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        return 2
    except KeyboardInterrupt:
        print("[ERROR] Đã dừng bởi người dùng.", file=sys.stderr)
        return 130


if __name__ == "__main__":
    raise SystemExit(main())
